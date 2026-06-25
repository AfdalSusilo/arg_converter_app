#!/usr/bin/env python3
"""
ARG Converter App
Konversi CSV logs menjadi Excel dan Word reports
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
import pandas as pd
import json
import os
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# =============================================================================
# CONFIGURATION & STYLING
# =============================================================================

SKY_BLUE = "#87CEEB"
DARK_NAVY = "#1E3A5F"
LIGHT_BLUE = "#E6F3FF"
WHITE = "#FFFFFF"
FONT_TITLE = ("Helvetica", 18, "bold")
FONT_LABEL = ("Helvetica", 12)
FONT_BUTTON = ("Helvetica", 11, "bold")


# =============================================================================
# DATA FLATTENING LOGIC
# =============================================================================

def flatten_json_column(df, column_name):
    """
    Flatten JSON string column into proper rows/columns.
    Returns a list of flattened records.
    """
    flattened_records = []

    for idx, row in df.iterrows():
        base_data = {col: row[col] for col in df.columns if col != column_name}
        json_str = row.get(column_name, None)

        if pd.isna(json_str) or json_str == "":
            flattened_records.append(base_data)
            continue

        try:
            json_data = json.loads(json_str)
            if isinstance(json_data, list):
                for item in json_data:
                    record = {**base_data, **item}
                    flattened_records.append(record)
            elif isinstance(json_data, dict):
                record = {**base_data, **json_data}
                flattened_records.append(record)
            else:
                flattened_records.append(base_data)
        except (json.JSONDecodeError, TypeError):
            flattened_records.append(base_data)

    return flattened_records


def detect_log_type(df):
    """
    Auto-detect log type based on column names.
    Returns: 'behavior', 'gui', 'npc_chat', or 'unknown'
    """
    columns_lower = [col.lower() for col in df.columns]

    if any(col in columns_lower for col in ['behavior', 'mouse_events', 'keystrokes']):
        return 'behavior'
    elif any(col in columns_lower for col in ['gui', 'input_data', 'widget']):
        return 'gui'
    elif any(col in columns_lower for col in ['chat', 'chat_history', 'npc_chat', 'message']):
        return 'npc_chat'
    return 'unknown'


def get_json_columns(df):
    """
    Return list of columns that contain JSON string data.
    """
    json_cols = []
    for col in df.columns:
        sample = df[col].dropna().head(5)
        if len(sample) > 0:
            try:
                for val in sample:
                    if isinstance(val, str) and val.strip().startswith(('[', '{')):
                        json.loads(val)
                        json_cols.append(col)
                        break
            except (json.JSONDecodeError, TypeError):
                continue
    return json_cols


def flatten_dataframe(df):
    """
    Flatten all JSON columns in the dataframe.
    """
    json_cols = get_json_columns(df)

    if not json_cols:
        return df.to_dict('records')

    # Flatten each JSON column one by one
    result_df = df.copy()
    for col in json_cols:
        flattened = flatten_json_column(result_df, col)
        result_df = pd.DataFrame(flattened)

    return result_df.to_dict('records')


# =============================================================================
# EXCEL EXPORT
# =============================================================================

def export_to_excel(df, output_path, sheet_name="ARG_Report"):
    """
    Export dataframe to formatted Excel file.
    - Auto-fit column widths
    - Text Wrap for long text
    - Sky Blue / Dark Navy header styling
    """

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    # Styles
    header_fill = PatternFill(start_color=DARK_NAVY, end_color=DARK_NAVY, fill_type="solid")
    header_font = Font(name='Calibri', bold=True, color=WHITE, size=11)
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    data_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Header row
    columns = list(df.columns)
    for col_idx, col_name in enumerate(columns, start=1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Data rows
    for row_idx, row_data in enumerate(df.itertuples(index=False), start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = data_alignment
            cell.border = thin_border

    # Auto-fit column widths
    for col_idx, col_name in enumerate(columns, start=1):
        max_length = len(str(col_name))
        for row_idx in range(2, len(df) + 2):
            cell_value = ws.cell(row=row_idx, column=col_idx).value
            if cell_value:
                max_length = max(max_length, len(str(cell_value)))

        adjusted_width = min(max_length + 2, 50)  # Cap at 50
        ws.column_dimensions[get_column_letter(col_idx)].width = adjusted_width

    # Freeze header row
    ws.freeze_panes = "A2"

    wb.save(output_path)


# =============================================================================
# WORD EXPORT
# =============================================================================

def export_to_word(df, output_path, log_type):
    """
    Export dataframe to formatted Word document.
    - Group by player_name (Heading 1)
    - Bold labels with italics/blockquotes for content
    """

    doc = Document()

    # Title
    title = doc.add_heading("ARG Log Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    meta_para.add_run(f"Log Type: {log_type.replace('_', ' ').title()}\n")
    meta_para.add_run(f"Total Records: {len(df)}")

    doc.add_paragraph()  # Spacer

    # Group by player_name if exists
    player_col = None
    for col in ['player_name', 'PlayerName', 'player', 'Player', 'username', 'Username']:
        if col in df.columns:
            player_col = col
            break

    if player_col:
        grouped = df.groupby(player_col)
        for player_name, group in grouped:
            # Heading 1 for each player
            heading = doc.add_heading(str(player_name), level=1)

            # Process each record
            for idx, row in group.iterrows():
                para = doc.add_paragraph()

                for col_name, value in row.items():
                    if col_name == player_col:
                        continue

                    if pd.isna(value) or value == "":
                        continue

                    # Label (bold)
                    label_run = para.add_run(f"{col_name}: ")
                    label_run.bold = True
                    label_run.font.size = Pt(11)

                    # Value (normal/italic for long text)
                    value_str = str(value)
                    value_run = para.add_run(value_str)
                    value_run.font.size = Pt(11)

                    para.add_run("\n")

                doc.add_paragraph()  # Spacer between records

    else:
        # No player grouping - just list all records
        for idx, row in df.iterrows():
            para = doc.add_paragraph()

            for col_name, value in row.items():
                if pd.isna(value) or value == "":
                    continue

                label_run = para.add_run(f"{col_name}: ")
                label_run.bold = True
                label_run.font.size = Pt(11)

                value_str = str(value)
                value_run = para.add_run(value_str)
                value_run.font.size = Pt(11)

                para.add_run("\n")

            doc.add_paragraph()

    doc.save(output_path)


# =============================================================================
# CUSTOMTKINTER UI
# =============================================================================

class ARGConverterApp(ctk.CTk):

    def __init__(self):
        super().__init__()

        # Window setup
        self.title("ARG Converter - CSV to Excel & Word")
        self.geometry("700x550")
        self.resizable(False, False)

        # Set appearance mode
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        # Variables
        self.csv_file_path = ctk.StringVar()
        self.log_type_var = ctk.StringVar(value="auto")
        self.df = None
        self.flattened_data = None

        self._build_ui()

    def _build_ui(self):
        """
        Build the customtkinter UI with Sky Blue theme.
        """

        # Main container with padding
        main_frame = ctk.CTkFrame(self, fg_color=LIGHT_BLUE, corner_radius=15)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Title
        title_label = ctk.CTkLabel(
            main_frame,
            text="ARG Converter",
            font=FONT_TITLE,
            text_color=DARK_NAVY
        )
        title_label.pack(pady=(10, 5))

        subtitle_label = ctk.CTkLabel(
            main_frame,
            text="Convert CSV Logs to Excel & Word Reports",
            font=("Helvetica", 10),
            text_color=DARK_NAVY
        )
        subtitle_label.pack(pady=(0, 20))

        # File selection frame
        file_frame = ctk.CTkFrame(main_frame, fg_color=WHITE, corner_radius=10)
        file_frame.pack(fill="x", padx=10, pady=10)

        # Select CSV button
        self.select_btn = ctk.CTkButton(
            file_frame,
            text="Select CSV File",
            command=self.select_csv_file,
            fg_color=SKY_BLUE,
            hover_color=DARK_NAVY,
            text_color=DARK_NAVY,
            font=FONT_BUTTON,
            corner_radius=8,
            height=40,
            width=200
        )
        self.select_btn.pack(pady=(15, 5), padx=15)

        # File path display
        self.file_path_label = ctk.CTkLabel(
            file_frame,
            text="No file selected",
            font=("Helvetica", 9),
            text_color="gray",
            wraplength=550
        )
        self.file_path_label.pack(pady=(0, 15), padx=15)

        # Log type selection
        log_type_frame = ctk.CTkFrame(main_frame, fg_color=WHITE, corner_radius=10)
        log_type_frame.pack(fill="x", padx=10, pady=10)

        log_type_title = ctk.CTkLabel(
            log_type_frame,
            text="Log Type Detection",
            font=FONT_LABEL,
            text_color=DARK_NAVY
        )
        log_type_title.pack(pady=(10, 5), padx=15)

        # Radio buttons
        radio_frame = ctk.CTkFrame(log_type_frame, fg_color="transparent")
        radio_frame.pack(pady=(0, 10), padx=15)

        self.radio_auto = ctk.CTkRadioButton(
            radio_frame,
            text="Auto-Detect",
            variable=self.log_type_var,
            value="auto",
            fg_color=SKY_BLUE,
            hover_color=DARK_NAVY
        )
        self.radio_auto.pack(side="left", padx=(0, 20))

        self.radio_behavior = ctk.CTkRadioButton(
            radio_frame,
            text="Behavior Logs",
            variable=self.log_type_var,
            value="behavior",
            fg_color=SKY_BLUE,
            hover_color=DARK_NAVY
        )
        self.radio_behavior.pack(side="left", padx=(0, 20))

        self.radio_gui = ctk.CTkRadioButton(
            radio_frame,
            text="GUI Logs",
            variable=self.log_type_var,
            value="gui",
            fg_color=SKY_BLUE,
            hover_color=DARK_NAVY
        )
        self.radio_gui.pack(side="left", padx=(0, 20))

        self.radio_npc = ctk.CTkRadioButton(
            radio_frame,
            text="NPC Chat",
            variable=self.log_type_var,
            value="npc_chat",
            fg_color=SKY_BLUE,
            hover_color=DARK_NAVY
        )
        self.radio_npc.pack(side="left")

        # Status label
        self.status_label = ctk.CTkLabel(
            log_type_frame,
            text="",
            font=("Helvetica", 10),
            text_color=DARK_NAVY
        )
        self.status_label.pack(pady=(0, 10), padx=15)

        # Export buttons frame
        export_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        export_frame.pack(fill="x", padx=10, pady=(10, 10))

        # Export to Excel button
        self.excel_btn = ctk.CTkButton(
            export_frame,
            text="Export to Excel",
            command=self.export_excel,
            fg_color=SKY_BLUE,
            hover_color=DARK_NAVY,
            text_color=DARK_NAVY,
            font=FONT_BUTTON,
            corner_radius=8,
            height=45,
            width=200
        )
        self.excel_btn.pack(side="left", padx=(5, 10))

        # Export to Word button
        self.word_btn = ctk.CTkButton(
            export_frame,
            text="Export to Word",
            command=self.export_word,
            fg_color=SKY_BLUE,
            hover_color=DARK_NAVY,
            text_color=DARK_NAVY,
            font=FONT_BUTTON,
            corner_radius=8,
            height=45,
            width=200
        )
        self.word_btn.pack(side="left", padx=(10, 5))

    def select_csv_file(self):
        """
        Open file dialog to select CSV file.
        """
        file_path = filedialog.askopenfilename(
            title="Select CSV File",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )

        if file_path:
            self.csv_file_path.set(file_path)
            self.file_path_label.configure(
                text=file_path,
                text_color=DARK_NAVY
            )

            try:
                # Load CSV
                self.df = pd.read_csv(file_path)
                self.flattened_data = flatten_dataframe(self.df)

                # Auto-detect log type
                detected_type = detect_log_type(self.df)
                self.log_type_var.set(detected_type)

                self.status_label.configure(
                    text=f"Loaded {len(self.df)} rows. Detected: {detected_type.replace('_', ' ').title()}"
                )

            except Exception as e:
                messagebox.showerror("Error", f"Failed to load CSV:\n{str(e)}")
                self.status_label.configure(text="Failed to load file")

    def export_excel(self):
        """
        Export to Excel with formatting.
        """
        if self.df is None:
            messagebox.showwarning("Warning", "Please select a CSV file first.")
            return

        output_path = filedialog.asksaveasfilename(
            title="Save Excel Report",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")]
        )

        if output_path:
            try:
                # Create flattened dataframe
                flattened_df = pd.DataFrame(self.flattened_data)

                export_to_excel(
                    flattened_df,
                    output_path,
                    sheet_name="ARG_Report"
                )

                messagebox.showinfo("Success", f"Excel report saved:\n{output_path}")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to export Excel:\n{str(e)}")

    def export_word(self):
        """
        Export to Word document.
        """
        if self.df is None:
            messagebox.showwarning("Warning", "Please select a CSV file first.")
            return

        output_path = filedialog.asksaveasfilename(
            title="Save Word Report",
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")]
        )

        if output_path:
            try:
                # Create flattened dataframe
                flattened_df = pd.DataFrame(self.flattened_data)

                log_type = self.log_type_var.get()
                if log_type == "auto":
                    log_type = detect_log_type(self.df)

                export_to_word(flattened_df, output_path, log_type)

                messagebox.showinfo("Success", f"Word report saved:\n{output_path}")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to export Word:\n{str(e)}")


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    app = ARGConverterApp()
    app.mainloop()
